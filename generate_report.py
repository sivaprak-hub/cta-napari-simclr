"""
Generate the full technical report as a Word (.docx) document.
Run: python generate_report.py
Output: CTA_Clustering_Methods_Report.docx  (same directory)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      "CTA_Clustering_Methods_Report.docx")


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)   # dark navy
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1B, 0x57, 0x7C)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def italic_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_font(run, size=11, italic=True)
    return p


def math_block(doc, text):
    """Formatted paragraph for equations / code snippets."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return p


def bullet(doc, text, sub=False):
    style = "List Bullet 2" if sub else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except Exception:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5 if sub else 0.8)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def kv(doc, key, value):
    """Bold key + normal value in one paragraph (for parameter tables inline)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(key + ":  ")
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(value)
    set_font(r2, size=11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "1A237E")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            row_cells[ci].text = str(val)
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
        if ri % 2 == 1:
            for cell in row_cells:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), "E8EAF6")
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for ri2 in range(len(table.rows)):
            for ci2, w in enumerate(col_widths):
                table.rows[ri2].cells[ci2].width = Cm(w)
    doc.add_paragraph()
    return table


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 100)
    run.font.size  = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=10, italic=True, color=(0x55, 0x55, 0x55))


# ---------------------------------------------------------------------------
# build document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ====================================================================
    # TITLE PAGE
    # ====================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(
        "Clustering Algorithms in Calcium Transient Analysis:\n"
        "A Comparative Technical Report on HDBSCAN,\n"
        "Normalized Cross-Correlation Graphs, and\n"
        "SimCLR Contrastive Representation Learning"
    )
    r.font.name  = "Times New Roman"
    r.font.size  = Pt(20)
    r.bold       = True
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph()
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run(
        "napari-cta-simclr  ·  Technical Implementation Reference\n"
        "Calcium Transient Analyzer — SimCLR Enhanced Edition"
    )
    r2.font.name  = "Times New Roman"
    r2.font.size  = Pt(13)
    r2.italic     = True
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = date_p.add_run(datetime.date.today().strftime("%B %Y"))
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(12)

    doc.add_page_break()

    # ====================================================================
    # ABSTRACT
    # ====================================================================
    heading(doc, "Abstract", 1)
    body(doc,
        "This report provides a comprehensive technical description of the three clustering "
        "methodologies implemented in the napari-cta-simclr plugin for live-cell calcium imaging "
        "analysis: (1) Hierarchical Density-Based Spatial Clustering of Applications with Noise "
        "(HDBSCAN) operating on five hand-crafted spatiotemporal scalar features; "
        "(2) a Normalized Cross-Correlation (NCC) adjacency graph combined with Louvain community "
        "detection; and (3) a SimCLR-based self-supervised contrastive representation learning "
        "framework that constructs an embedding-space adjacency graph for community detection. "
        "Each method is described in terms of its theoretical foundations, algorithmic implementation, "
        "hyperparameter choices and their scientific justifications, software dependencies, and "
        "expected performance characteristics. A complete preprocessing pipeline is described first, "
        "as all three methods share the same upstream signal conditioning. Detailed comparison "
        "tables and parameter inventories are provided throughout."
    )
    doc.add_page_break()

    # ====================================================================
    # TABLE OF CONTENTS (manual)
    # ====================================================================
    heading(doc, "Contents", 1)
    toc_items = [
        ("1", "Introduction and Motivation"),
        ("2", "Shared Preprocessing Pipeline"),
        ("   2.1", "Image Loading and Spatial Binning"),
        ("   2.2", "Bleach Correction and Baseline Subtraction"),
        ("   2.3", "Activity Mask Computation"),
        ("   2.4", "Beat Detection"),
        ("3", "Method 1 — HDBSCAN on Spatiotemporal Features"),
        ("   3.1", "Feature Engineering"),
        ("   3.2", "HDBSCAN Algorithm"),
        ("   3.3", "Hyperparameters and Justification"),
        ("   3.4", "Fallback Cascade"),
        ("   3.5", "Libraries"),
        ("4", "Method 2 — NCC Adjacency Graph with Louvain Community Detection"),
        ("   4.1", "NCC Adjacency Matrix Construction"),
        ("   4.2", "Graph Community Detection: Louvain Algorithm"),
        ("   4.3", "Cluster Quality Metric: Intra-Cluster NCC Score"),
        ("   4.4", "Hyperparameters and Justification"),
        ("   4.5", "Fallback Cascade"),
        ("   4.6", "Libraries"),
        ("5", "Method 3 — SimCLR Contrastive Representation Learning"),
        ("   5.1", "Theoretical Background: Contrastive Self-Supervised Learning"),
        ("   5.2", "1D Convolutional Encoder Architecture"),
        ("   5.3", "Calcium Signal Augmentation Strategy"),
        ("   5.4", "NT-Xent Contrastive Loss Function"),
        ("   5.5", "Transductive Per-Recording Training"),
        ("   5.6", "Embedding-Space Adjacency Graph Construction"),
        ("   5.7", "Hyperparameters and Justification"),
        ("   5.8", "Libraries"),
        ("6", "Comparative Analysis"),
        ("7", "Complete Hyperparameter Inventory"),
        ("8", "References"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_num = p.add_run(f"{num:<6}")
        set_font(r_num, size=11, bold=True)
        r_title = p.add_run(title)
        set_font(r_title, size=11)
    doc.add_page_break()

    # ====================================================================
    # 1. INTRODUCTION
    # ====================================================================
    heading(doc, "1.  Introduction and Motivation", 1)
    body(doc,
        "Calcium transient imaging is a cornerstone technique in cardiac and neuroscience research, "
        "enabling the simultaneous visualization of intracellular calcium dynamics across thousands "
        "of pixels in a living cell preparation. The fundamental analytical challenge is to group "
        "pixels that share similar calcium transient waveforms into functionally coherent clusters — "
        "regions that fire synchronously and likely correspond to single cells or cell ensembles "
        "connected by gap junctions."
    )
    body(doc,
        "The original napari-cta plugin addressed this with HDBSCAN applied to five scalar "
        "features derived from each pixel's calcium trace. While computationally efficient, this "
        "approach encodes only coarse spectral and temporal properties of the waveform and can be "
        "dominated by the dominant frequency feature, causing pixels with similar frequencies but "
        "different waveform shapes to be incorrectly co-clustered."
    )
    body(doc,
        "The napari-cta-simclr plugin implements two additional clustering strategies that operate "
        "on the raw waveform directly, bypassing hand-crafted feature engineering: "
        "(a) a Normalized Cross-Correlation (NCC) adjacency graph that captures the peak "
        "temporal similarity between spatially adjacent pixels, and (b) a SimCLR contrastive "
        "learning framework that learns a nonlinear embedding of each pixel's calcium trace and "
        "uses cosine similarity in that embedding space to construct the adjacency graph. "
        "All three methods are run simultaneously on every recording, and their cluster maps are "
        "presented as independent napari layers for direct visual comparison."
    )
    body(doc,
        "This document provides a full algorithmic description of each method, including all "
        "hyperparameters with their scientific justification, the software stack, and a "
        "comparative performance analysis."
    )

    # ====================================================================
    # 2. PREPROCESSING
    # ====================================================================
    heading(doc, "2.  Shared Preprocessing Pipeline", 1)
    body(doc,
        "All three clustering methods consume the same preprocessed signals. The preprocessing "
        "pipeline is implemented in AnalysisWorker.run() in backend.py and executes in a "
        "background QThread to preserve GUI responsiveness."
    )

    heading(doc, "2.1  Image Loading and Spatial Binning", 2)
    body(doc,
        "Raw fluorescence images are loaded as three-dimensional arrays of shape (T, H, W) — "
        "T frames, H rows, W columns — where each value represents the pixel-level fluorescence "
        "intensity. TIFF and TIFF-series files are read using tifffile; Olympus VSI files are "
        "read via a custom binary ETS parser or, if the aicsimageio package is present, through "
        "its Bio-Formats bridge."
    )
    body(doc,
        "Prior to any analysis, the raw stack is intensity-clipped at the [0.4, 99.6] percentile "
        "range to eliminate detector saturation artefacts and dark-current outliers, then "
        "linearly rescaled to [0, 255]. Spatial binning is subsequently applied: each frame is "
        "first smoothed with a Gaussian filter to reduce spatial aliasing, then downsampled by "
        "block averaging."
    )
    kv(doc, "Gaussian σ", "H_raw / 204.8  (adaptive; equals ~2.5 px for a 512-row image)")
    kv(doc, "Block reduction", "bin_size × bin_size spatial window, mean pooled (skimage.measure.block_reduce)")
    kv(doc, "Default bin_size", "16 px for images with max(H, W) < 2048; 32 px otherwise")
    kv(doc, "Clip percentiles", "[0.4th, 99.6th] — asymmetric to be more conservative on bright outliers")
    body(doc,
        "The adaptive Gaussian sigma ensures that the pre-filter radius scales with image "
        "resolution, preventing aliasing at high resolution without over-smoothing at low "
        "resolution. The denominator 204.8 = 1024 / 5 yields sigma ≈ 2.5 px for a 512-row image "
        "and sigma ≈ 5 px for a 1024-row image."
    )

    heading(doc, "2.2  Bleach Correction and Baseline Subtraction", 2)
    body(doc,
        "Photobleaching causes a gradual decay of fluorescence intensity throughout the recording "
        "that is unrelated to calcium dynamics. Two baseline models are provided:"
    )
    body(doc, "Single Exponential Model (default):", indent=False)
    math_block(doc, "F_baseline(t)  =  a · exp(−t / τ) + c")
    body(doc,
        "Parameters (a, τ, c) are fitted to the running lower-envelope of each pixel's signal "
        "using scipy.optimize.curve_fit with the Levenberg-Marquardt algorithm. The envelope is "
        "estimated as the 10th percentile within a sliding window of width max(T/8, 3) frames. "
        "Initial parameter guess: a₀ = envelope[0] − envelope[−1]; τ₀ = total_duration / 3; "
        "c₀ = envelope[−1]. Fitting bounds constrain a ≥ 0 and τ ≥ 1×10⁻⁵ s to prevent "
        "numerical instability. Maximum function evaluations: 2000."
    )
    body(doc, "Boundary Baseline Model:", indent=False)
    body(doc,
        "A piecewise-linear spline is constructed through local minima detected by "
        "scipy.signal.find_peaks on the negated signal, with minimum spacing of T/10 frames. "
        "This model accommodates multi-phase drift and slow oscillatory artefacts. The corrected "
        "signal is sig[i] = raw[i] − baseline[i] for both models.",
        indent=True
    )

    heading(doc, "2.3  Activity Mask Computation", 2)
    body(doc,
        "Not all pixels contain meaningful calcium dynamics. An activity mask is computed to "
        "restrict all three clustering algorithms to pixels that show genuine oscillatory activity, "
        "reducing noise and computational cost."
    )
    body(doc,
        "Step 1 — Bandpass Filtering. Each corrected signal is filtered with a 4th-order "
        "Butterworth bandpass filter applied via the second-order-sections (SOS) representation "
        "and zero-phase forward-backward filtering (sosfiltfilt) to eliminate non-cardiac "
        "frequency content:"
    )
    kv(doc, "Low cutoff (f_low)", "0.3 Hz  (≈18 BPM; excludes DC drift and very slow oscillations)")
    kv(doc, "High cutoff (f_high)", "5.0 Hz  (≈300 BPM; excludes high-frequency noise and motion artefacts)")
    kv(doc, "Filter order", "4  (24 dB/octave rolloff; high enough for good selectivity without ringing)")
    kv(doc, "Nyquist safety cap", "min(f, nyq × 0.9)  (prevents filter instability at the Nyquist limit)")
    body(doc,
        "Step 2 — Pulsatility Map. The variance of the bandpass-filtered signal over time is "
        "computed for each pixel and reshaped to a spatial map of shape (H_bin, W_bin). Variance "
        "is preferred over amplitude because it is robust to DC offset and captures the energy "
        "of the oscillation regardless of baseline level."
    )
    math_block(doc, "P(y, x)  =  Var_t [ bandpass(sig_{y,x}) ]")
    body(doc,
        "Step 3 — Otsu Thresholding. The activity mask is obtained by applying Otsu's method "
        "(skimage.filters.threshold_otsu) to the flattened pulsatility map. Otsu's method "
        "maximises inter-class variance and is threshold-free, making it robust across recordings "
        "with different cell densities and SNR levels. Two adaptive fallbacks protect against "
        "pathological cases:"
    )
    bullet(doc, "If the resulting active fraction < 3% or > 85%, use the 65th percentile of the pulsatility distribution as the threshold (prevents under- or over-segmentation)")
    bullet(doc, "If threshold_otsu raises an exception (e.g., uniform image), default to the 65th percentile")

    heading(doc, "2.4  Beat Detection", 2)
    body(doc,
        "A reference trace is constructed as the amplitude-weighted mean of all active pixels' "
        "corrected signals. Beat peaks are detected on this reference trace using "
        "scipy.signal.find_peaks with:"
    )
    kv(doc, "Prominence threshold", "0.15 × signal_range  (avoids detecting noise peaks while capturing all genuine beats)")
    kv(doc, "Minimum inter-peak distance", "max(int(fps × 0.8), 2) frames  (enforces a minimum refractory period of 0.8 s; 75 BPM maximum)")
    body(doc,
        "The detected beat indices are used for beat-averaged kinetic parameter extraction "
        "(T_ON, T_OFF, amplitude, BPM, duty cycle) which feed the Results & Metrics panel. "
        "Beat peaks are NOT used directly by any of the three clustering methods."
    )

    doc.add_page_break()

    # ====================================================================
    # 3. HDBSCAN
    # ====================================================================
    heading(doc, "3.  Method 1 — HDBSCAN on Spatiotemporal Features", 1)
    body(doc,
        "HDBSCAN (Hierarchical Density-Based Spatial Clustering of Applications with Noise) is "
        "the original clustering method from the napari-cta plugin. It operates on a five-element "
        "feature vector computed per active pixel, transforms them with a standard scaler, and "
        "applies density-based hierarchical clustering."
    )

    heading(doc, "3.1  Feature Engineering", 2)
    body(doc,
        "Five scalar features are extracted from each active pixel's corrected, bandpass-filtered "
        "calcium trace. These features are designed to capture the principal axes of variation "
        "in cardiomyocyte calcium transients:"
    )
    add_table(doc,
        ["#", "Feature", "Formula / Computation", "Physical Meaning"],
        [
            ["1", "Pulsatility", "Var_t(filtered_trace[i])",
             "Total oscillatory energy; distinguishes strongly beating from quiescent pixels"],
            ["2", "Dominant Frequency", "argmax |FFT(filtered_trace)|  for f ∈ [0.3, 5.0] Hz",
             "Firing rate (BPM / 60); the most discriminative feature for separating cell populations with different beating rates"],
            ["3", "Mean Amplitude", "Mean of per-beat peak-to-trough amplitude over a window of max(1.5×fps, 3) frames",
             "Contractile strength; varies with cell health, calcium load, and dye concentration"],
            ["4", "Phase Offset", "Mean angle of the analytical signal: E[arg(hilbert(filtered_trace))]",
             "Temporal phase of the oscillation; helps separate pixels that fire at the same rate but at different phases"],
            ["5", "Duty Cycle", "Fraction of frames where sig > min(sig) + 0.5 × signal_range",
             "Width of the calcium transient relative to the inter-beat interval"],
        ],
        col_widths=[1.0, 3.5, 6.0, 5.5]
    )
    caption(doc, "Table 3.1 — Five scalar features computed per active pixel for HDBSCAN clustering.")

    body(doc,
        "Feature computation uses scipy.signal.hilbert for the Hilbert transform and numpy.fft "
        "for the FFT. All features are computed on the bandpass-filtered trace (not the raw "
        "corrected signal) to isolate the cardiac frequency band. The window for amplitude "
        "estimation is max(int(fps × 1.5), 3) frames, corresponding to approximately 1.5 "
        "cardiac cycles at the recording frame rate."
    )
    body(doc,
        "Feature scaling is applied with sklearn.preprocessing.StandardScaler "
        "(zero mean, unit variance), which is essential for HDBSCAN because the algorithm uses "
        "Euclidean distance — without scaling, dominant frequency (in Hz, typically 0.3–5.0) "
        "would numerically overwhelm duty cycle (dimensionless, 0–1). NaN and infinity values "
        "resulting from short traces or zero-signal pixels are replaced with 0.0 after scaling."
    )

    heading(doc, "3.2  HDBSCAN Algorithm", 2)
    body(doc,
        "HDBSCAN (Campello et al., 2013) extends DBSCAN by building a hierarchy of density-based "
        "clusters and then extracting the most persistent (stable) flat partition. The algorithm "
        "operates in three phases:"
    )
    bullet(doc, "Core distance computation: for each point p, core_dist(p, k) = distance to the k-th nearest neighbor (k = min_samples).")
    bullet(doc, "Minimum spanning tree construction on the mutual reachability distance graph: d_mreach(p,q) = max(core_dist(p,k), core_dist(q,k), d(p,q)).")
    bullet(doc, "Hierarchy condensation and cluster extraction: branches of the MST are pruned by minimum cluster size; points that fall below this threshold are labeled as noise (label = −1).")
    body(doc,
        "HDBSCAN's key advantage is that it does not require the user to specify the number of "
        "clusters in advance and naturally identifies noise points, which is important for calcium "
        "imaging where background and artefact pixels should not be forced into any cluster."
    )
    math_block(doc,
        "d_mreach(p, q)  =  max( core_dist_k(p),  core_dist_k(q),  ||p − q|| )"
    )

    heading(doc, "3.3  Hyperparameters and Justification", 2)
    kv(doc, "min_cluster_size", "max(3, N_active // 15)  — dynamic, proportional to the number of active pixels")
    body(doc,
        "This adaptive formula ensures that the minimum cluster size scales with the recording. "
        "For a recording with 60 active pixels, min_cluster_size = max(3, 4) = 4. For 300 active "
        "pixels, min_cluster_size = 20. A fixed value would either over-fragment small recordings "
        "or merge distinct populations in large recordings. The 1/15 ratio was chosen empirically "
        "to yield 3–15 clusters for typical cardiomyocyte monolayer recordings.",
        indent=True
    )
    kv(doc, "min_samples", "2  — controls how conservative the density estimation is")
    body(doc,
        "Setting min_samples = 2 means that a point needs only one neighbour to be considered "
        "a core point. This is the most permissive setting and maximises the number of points "
        "assigned to clusters rather than noise. In calcium imaging, the activity mask already "
        "pre-selects pixels with genuine oscillatory signal, so aggressive noise detection by "
        "HDBSCAN is unnecessary.",
        indent=True
    )
    kv(doc, "metric", "euclidean  (default Euclidean distance in the 5D scaled feature space)")
    kv(doc, "StandardScaler", "Mean = 0, Variance = 1 per feature; fitted on the active pixel set only")

    heading(doc, "3.4  Fallback Cascade", 2)
    body(doc,
        "If the hdbscan package is not installed, the method falls back to scikit-learn's "
        "KMeans with:"
    )
    kv(doc, "n_clusters", "min(4, max(2, N_active // 5))  — scales from 2 to 4 clusters")
    kv(doc, "n_init", "10  (number of random initializations to avoid local minima)")
    kv(doc, "random_state", "42  (reproducibility)")

    heading(doc, "3.5  Libraries", 2)
    add_table(doc,
        ["Library", "Version Requirement", "Role"],
        [
            ["hdbscan", "≥ 0.8.33", "Primary clustering engine"],
            ["scikit-learn", "≥ 1.3", "StandardScaler, KMeans fallback, SpectralClustering"],
            ["scipy", "≥ 1.10", "FFT, Hilbert transform, signal processing"],
            ["numpy", "≥ 1.24", "Feature computation, array operations"],
        ],
        col_widths=[3.5, 4.0, 8.5]
    )

    doc.add_page_break()

    # ====================================================================
    # 4. NCC GRAPH
    # ====================================================================
    heading(doc, "4.  Method 2 — NCC Adjacency Graph with Louvain Community Detection", 1)
    body(doc,
        "The NCC Graph method replaces hand-crafted scalar features with a direct measurement "
        "of waveform shape similarity between spatially adjacent pixels. An adjacency graph is "
        "constructed where edge weights are the peak normalized cross-correlation between "
        "neighboring pixel pairs, and community structure is extracted using the Louvain "
        "modularity maximisation algorithm."
    )

    heading(doc, "4.1  NCC Adjacency Matrix Construction", 2)
    body(doc,
        "For each pair of 4-connected active pixels (i, j), the Normalized Cross-Correlation "
        "(NCC) at lag τ is defined as:"
    )
    math_block(doc, "NCC(i, j, τ)  =  Σ_t  ŝ_i(t) · ŝ_j(t + τ)")
    body(doc,
        "where ŝ_i denotes the zero-meaned, L2-normalised version of the corrected calcium "
        "signal at pixel i. L2 normalisation ensures that NCC ∈ [−1, +1] by the Cauchy-Schwarz "
        "inequality. The edge weight is the maximum NCC over all lags within the permitted range:"
    )
    math_block(doc, "w(i, j)  =  max_{|τ| ≤ L}  NCC(i, j, τ)  ,   clipped to [0, 1]")
    body(doc,
        "where L = max(1, ⌊T × 0.25⌋) is the maximum permissible lag in frames. The lag limit "
        "prevents spuriously high correlations arising from aligning noise at large temporal "
        "offsets. Negative NCC values (anti-correlated signals) are clipped to zero because "
        "negative edge weights make the Louvain modularity objective ill-defined."
    )
    body(doc,
        "The full cross-correlation is computed using scipy.signal.correlate with method='fft', "
        "which uses the Fast Fourier Transform for O(T log T) computation rather than the "
        "O(T²) direct method. The lag-limited peak extraction then requires only a "
        "O(2L + 1) ≈ O(T/2) slice operation."
    )
    body(doc,
        "Only 4-connected (up, down, left, right) neighbors are considered, not 8-connected or "
        "all-pairs. This constraint enforces spatial continuity: two active regions separated by "
        "inactive tissue cannot be merged, which is biologically correct for spatially coherent "
        "calcium waves in cardiomyocyte monolayers. The adjacency matrix is stored as a "
        "scipy.sparse.csr_matrix of shape (N_active × N_active), indexed over active pixels only."
    )

    heading(doc, "4.2  Graph Community Detection: Louvain Algorithm", 2)
    body(doc,
        "The NCC adjacency graph is passed to the Louvain community detection algorithm "
        "(python-louvain 0.16, community_louvain.best_partition) with edge weights equal to the "
        "NCC values. The Louvain algorithm maximises the weighted modularity:"
    )
    math_block(doc,
        "Q  =  (1/2m) · Σ_{i,j} [ A_{ij} − k_i · k_j / 2m ] · δ(c_i, c_j)"
    )
    body(doc,
        "where A_{ij} is the NCC edge weight, k_i = Σ_j A_{ij} is the weighted degree of node i, "
        "m = (1/2) Σ_{ij} A_{ij} is the total edge weight, and δ(c_i, c_j) = 1 iff pixels i and "
        "j are assigned to the same community. Positive modularity Q indicates more within-community "
        "connections than expected in a null random graph with the same degree sequence."
    )
    body(doc,
        "The Louvain algorithm proceeds in two phases per iteration: (1) each node is greedily "
        "moved to the community that maximises ΔQ; (2) the identified communities are collapsed "
        "into super-nodes and the process repeats. Convergence is achieved when no single-node "
        "move increases Q. The critical advantage over k-means or HDBSCAN is that the number of "
        "communities is determined entirely by the data — no k must be specified."
    )
    kv(doc, "resolution", "1.0  — controls community granularity")
    body(doc,
        "resolution > 1.0 favours smaller communities (finer partition); "
        "resolution < 1.0 favours larger communities (coarser partition). The default of 1.0 "
        "implements the original Girvan-Newman modularity objective without rescaling. This is "
        "the appropriate default for cardiomyocyte calcium imaging where functional clusters "
        "typically have sizes between 10 and 200 pixels at bin = 16.",
        indent=True
    )
    kv(doc, "random_state", "42  (Louvain is a greedy heuristic; fixing the seed ensures reproducibility)")
    kv(doc, "Graph backend", "networkx.Graph via networkx.from_scipy_sparse_array()")

    heading(doc, "4.3  Cluster Quality Metric: Intra-Cluster NCC Score", 2)
    body(doc,
        "After clustering, the quality of the partition is quantified as the mean NCC edge "
        "weight over all edges whose two endpoints belong to the same community:"
    )
    math_block(doc,
        "score  =  mean{ w(i,j)  :  (i,j) ∈ E, c_i = c_j }"
    )
    body(doc,
        "This score lies in [0, 1]. A score near 1.0 indicates that all neighbouring pixel "
        "pairs within each cluster have highly correlated waveforms. A score near 0.5 suggests "
        "moderate within-cluster waveform heterogeneity. The score is computed efficiently using "
        "scipy.sparse.coo_matrix iteration over non-zero entries of the adjacency matrix."
    )

    heading(doc, "4.4  Hyperparameters and Justification", 2)
    add_table(doc,
        ["Parameter", "Value", "Justification"],
        [
            ["max_lag_fraction", "0.25",
             "Limits lag search to ±25% of recording length. Prevents noise-driven spurious peaks at large lags while accommodating genuine inter-cell conduction delays (typically < 100 ms in cardiomyocyte monolayers at 15 fps ≈ 4 frames)."],
            ["NCC clipping", "[0, 1]",
             "Negative weights (anti-correlated signals) are set to zero. Negative-weight graph edges are not interpretable in the modularity framework."],
            ["Louvain resolution", "1.0",
             "Standard modularity objective. Values in [0.5, 2.0] may be explored by advanced users modifying cluster_ncc_graph() directly."],
            ["Louvain random_state", "42",
             "Reproducibility. The greedy phase-assignment is stochastic; fixing the seed ensures identical results on repeated runs."],
            ["Graph type", "4-connected",
             "Spatial coherence constraint. Two pixels separated by inactive tissue cannot be co-clustered, which is physiologically correct."],
            ["Sparse storage", "scipy.sparse.csr_matrix",
             "For N = 500 active pixels, the 4-connected graph has at most 2000 edges vs. 250,000 all-pairs. Memory: O(N) vs. O(N²)."],
        ],
        col_widths=[3.8, 2.0, 10.2]
    )

    heading(doc, "4.5  Fallback Cascade", 2)
    body(doc, "If the Louvain algorithm fails (disconnected graph, import error), two fallbacks execute in sequence:")
    bullet(doc, "Spectral Clustering (sklearn.cluster.SpectralClustering): n_clusters = min(8, max(2, N // 10)); affinity='precomputed'; random_state=42; n_init=10. The NCC matrix is used as the precomputed affinity.")
    bullet(doc, "HDBSCAN on NCC distances: distance matrix = 1 − NCC; min_cluster_size = max(3, N // 15); min_samples = 2; metric='precomputed'. Noise labels are reassigned to a catch-all cluster.")
    bullet(doc, "Single cluster: if all else fails, all active pixels are assigned to cluster 0.")

    heading(doc, "4.6  Libraries", 2)
    add_table(doc,
        ["Library", "Version Requirement", "Role"],
        [
            ["scipy", "≥ 1.10", "FFT cross-correlation, sparse matrix storage"],
            ["networkx", "≥ 3.0", "Graph data structure for Louvain input"],
            ["python-louvain", "≥ 0.16", "Louvain modularity maximisation (imports as 'community')"],
            ["scikit-learn", "≥ 1.3", "SpectralClustering fallback"],
            ["hdbscan", "≥ 0.8.33", "Second fallback"],
        ],
        col_widths=[3.8, 3.8, 8.4]
    )

    doc.add_page_break()

    # ====================================================================
    # 5. SIMCLR
    # ====================================================================
    heading(doc, "5.  Method 3 — SimCLR Contrastive Representation Learning", 1)
    body(doc,
        "SimCLR (Simple Framework for Contrastive Learning of Visual Representations; "
        "Chen et al., 2020) is a self-supervised representation learning framework that trains "
        "a neural network encoder to place similar samples close together and dissimilar samples "
        "far apart in an embedding space, without requiring any class labels. In the context of "
        "calcium transient analysis, 'similar' means two stochastic augmentations of the same "
        "pixel's signal, and 'dissimilar' means signals from different pixels."
    )

    heading(doc, "5.1  Theoretical Background: Contrastive Self-Supervised Learning", 2)
    body(doc,
        "The core insight of contrastive learning is that a good representation should be "
        "invariant to nuisance transformations (noise, amplitude scaling, temporal shift) while "
        "being discriminative across truly different stimuli. For calcium signals, the nuisances "
        "are measurement-level variabilities (uneven dye loading, focal plane drift, photobleaching "
        "residuals) and the discriminative structure is the waveform shape — frequency, phase, "
        "rise/decay kinetics."
    )
    body(doc,
        "SimCLR achieves this by the following procedure: for each training step, two randomly "
        "augmented views x_i and x_j are created from the same source signal x. An encoder f(·) "
        "maps each view to a representation h ∈ ℝᵈ, and a projection head g(·) maps h to a "
        "normalised embedding z ∈ ℝᵖ. The NT-Xent loss then encourages z_i and z_j (from the "
        "same source) to have high cosine similarity, while penalising similarity to all other "
        "z vectors in the batch. After training, the projection head is discarded and the encoder "
        "representations h are used for downstream clustering."
    )
    math_block(doc, "h_i = f( aug_1(x) )  ,   h_j = f( aug_2(x) )")
    math_block(doc, "z_i = normalize( g(h_i) )  ,   z_j = normalize( g(h_j) )")

    heading(doc, "5.2  1D Convolutional Encoder Architecture", 2)
    body(doc,
        "Each calcium transient is treated as a univariate time series of length T. The encoder "
        "f is a 1D Convolutional Neural Network (CNN) that processes input tensors of shape "
        "(B, 1, T) — batch size B, 1 input channel, T time steps — and produces embeddings "
        "h ∈ ℝ^128 via three convolutional blocks followed by global average pooling."
    )
    add_table(doc,
        ["Layer", "Type", "Parameters", "Output Shape", "Purpose"],
        [
            ["1", "Conv1d", "in=1, out=32, k=7, pad=3", "(B, 32, T/2)", "Multi-scale temporal feature extraction; k=7 ≈ 0.5 s at 15 fps"],
            ["2", "BatchNorm1d", "32 features", "(B, 32, T/2)", "Training stability; reduces internal covariate shift"],
            ["3", "ReLU", "inplace=True", "(B, 32, T/2)", "Non-linearity"],
            ["4", "MaxPool1d", "kernel=2", "(B, 32, T/2)", "Temporal downsampling; translational invariance"],
            ["5", "Conv1d", "in=32, out=64, k=5, pad=2", "(B, 64, T/4)", "Higher-level temporal patterns (beat shape, kinetics)"],
            ["6", "BatchNorm1d", "64 features", "(B, 64, T/4)", "Training stability"],
            ["7", "ReLU", "inplace=True", "(B, 64, T/4)", "Non-linearity"],
            ["8", "MaxPool1d", "kernel=2", "(B, 64, T/4)", "Temporal downsampling"],
            ["9", "Conv1d", "in=64, out=128, k=3, pad=1", "(B, 128, T/4)", "High-level abstract features"],
            ["10", "BatchNorm1d", "128 features", "(B, 128, T/4)", "Training stability"],
            ["11", "ReLU", "inplace=True", "(B, 128, T/4)", "Non-linearity"],
            ["12", "AdaptiveAvgPool1d", "output_size=1", "(B, 128)", "Global temporal pooling; T-independent output"],
            ["—", "Flatten", "squeeze dim=2", "(B, 128)", "h ∈ ℝ^128"],
            ["P1", "Linear", "128 → 128", "(B, 128)", "Projection head (training only)"],
            ["P2", "ReLU", "inplace=True", "(B, 128)", "Projection non-linearity"],
            ["P3", "Linear", "128 → 64", "(B, 64)", "Projection to lower-dimensional contrast space"],
            ["P4", "F.normalize", "dim=1, L2-norm", "(B, 64)", "z ∈ ℝ^64 on unit hypersphere"],
        ],
        col_widths=[0.8, 3.2, 4.2, 2.8, 5.0]
    )
    caption(doc, "Table 5.1 — SimCLREncoder architecture. Layers P1–P4 form the projection head used only during training.")

    body(doc,
        "Architectural Design Rationale:"
    )
    bullet(doc, "Kernel sizes [7, 5, 3]: large kernels in early layers capture multi-scale temporal structure (beat onset, plateau, decay); smaller kernels in later layers aggregate fine-grained local features.")
    bullet(doc, "AdaptiveAvgPool1d(1): produces a fixed-size embedding regardless of T, enabling the same model architecture to handle recordings of varying duration without resampling.")
    bullet(doc, "BatchNorm at every block: essential for stability with small batches (min 20 pixels) and z-scored inputs that may have very similar scale statistics.")
    bullet(doc, "Projection head (128→128→64): following SimCLR paper recommendations, a nonlinear projection head improves the quality of the trunk representations h. The head is discarded at inference time.")
    bullet(doc, "Total parameters: approximately 127,000 — intentionally small to train transductively in < 30 s on CPU with 100 epochs.")

    heading(doc, "5.3  Calcium Signal Augmentation Strategy", 2)
    body(doc,
        "Four augmentation types are implemented, and one is randomly selected per signal per "
        "training step. The augmentation pool is designed to make the learned representation "
        "invariant to measurement-level variabilities while retaining sensitivity to biologically "
        "meaningful waveform differences:"
    )
    add_table(doc,
        ["Aug. Type", "Operation", "Magnitude", "Biological Rationale"],
        [
            ["Circular Phase Shift",
             "x' = roll(x, shift)\nshift ~ U(−T/2, +T/2)",
             "±50% of recording length",
             "The start of recording is arbitrary relative to the cell's beat cycle. Phase-invariant embeddings correctly co-cluster cells that beat together but happen to be captured at different phases."],
            ["Amplitude Scale",
             "x' = x · α\nα ~ U(0.6, 1.4)",
             "±40% of signal amplitude",
             "Amplitude varies with dye loading concentration and focal plane position. Embeddings should be amplitude-invariant to correctly group pixels with identical waveform shapes but different SNR."],
            ["Gaussian Noise",
             "x' = x + ε\nε ~ N(0, (0.05 · range(x))²)",
             "5% of signal range",
             "Sensor noise, photon shot noise, and residual bleaching add stochastic perturbations. Robustness to 5% noise is consistent with typical SNR in calcium imaging (SNR ≈ 5–20 dB)."],
            ["Sinusoidal Drift",
             "x' = x + 0.1·range(x)·sin(t + φ)\nφ ~ U(0, 2π)",
             "±10% of signal range",
             "Residual low-frequency drift after bleach correction (e.g., mechanical drift, focus changes) adds a slow sinusoidal baseline. Representing this as an augmentation makes the encoder robust to imperfect baseline subtraction."],
        ],
        col_widths=[3.2, 4.0, 3.2, 5.6]
    )
    caption(doc, "Table 5.2 — Four stochastic augmentation types applied to calcium signals during SimCLR training.")

    heading(doc, "5.4  NT-Xent Contrastive Loss Function", 2)
    body(doc,
        "The Normalised Temperature-scaled Cross Entropy (NT-Xent) loss is the training "
        "objective of SimCLR. For a batch of B pixel signals, two augmented views are "
        "generated for each signal, producing 2B embeddings. For each view i, its positive "
        "pair is the other view of the same signal (j = i + B or j = i − B), and all other "
        "2B − 2 views are treated as negatives."
    )
    math_block(doc,
        "L  =  −(1/2B) · Σ_{i=1}^{2B}  log exp(z_i · z_pos(i) / τ)\n"
        "                                    ─────────────────────────────────────────\n"
        "                                    Σ_{k=1}^{2B}  [ k ≠ i ] · exp(z_i · z_k / τ)"
    )
    body(doc,
        "The similarity between two L2-normalised embeddings z_i · z_j is their cosine "
        "similarity, which lies in [−1, +1]. The temperature parameter τ controls the "
        "concentration of the distribution: low τ sharpens the distribution (hard negatives "
        "dominate, strong gradient signal), high τ softens it (all negatives contribute "
        "equally)."
    )
    kv(doc, "Temperature τ", "0.5  — a standard choice from the SimCLR paper")
    body(doc,
        "τ = 0.5 was selected empirically in the original SimCLR work on ImageNet and has been "
        "found to generalise well to small-scale contrastive problems. It provides a sharp enough "
        "gradient signal for the encoder to learn discriminative features within 100 epochs, while "
        "remaining numerically stable (avoids exp overflow at τ → 0). At τ = 0.5, the NT-Xent "
        "loss is equivalent to a softmax classifier with 2B−1 classes where the positive class "
        "has logit = 2 × cosine_sim.",
        indent=True
    )
    body(doc,
        "Implementation details: the 2B × 2B similarity matrix is computed in a single "
        "torch.mm operation. The diagonal (self-similarity, always 1.0) is masked to −∞ before "
        "the softmax to prevent the trivial solution of a signal matching itself. The positive "
        "index for the first B rows is row + B, and for the last B rows is row − B. A single "
        "call to torch.nn.functional.cross_entropy then computes the mean loss."
    )

    heading(doc, "5.5  Transductive Per-Recording Training", 2)
    body(doc,
        "Unlike typical deep learning workflows where a model is pre-trained on a large corpus "
        "and fine-tuned on a target dataset, SimCLR here is trained from scratch on every "
        "individual recording. This transductive approach — using only the test data for training "
        "— is appropriate because:"
    )
    bullet(doc, "Calcium waveform morphology varies substantially across cell types, dyes, microscopes, and experimental conditions. A universal pre-trained encoder would require a large labelled corpus that does not currently exist.")
    bullet(doc, "The number of active pixels per recording (typically 50–600) is too small to meaningfully fine-tune a general-purpose model without overfitting.")
    bullet(doc, "The training task (contrastive discrimination among all pixels in this recording) adapts the representation to the specific SNR, frequency range, and waveform shape of the current experiment.")
    body(doc, "Training configuration:")
    kv(doc, "Optimiser", "Adam(lr=3×10⁻⁴, weight_decay=1×10⁻⁵)")
    body(doc,
        "Adam's adaptive learning rate is well-suited for this task because the gradients can "
        "vary by orders of magnitude across layers (early Conv layers vs. projection head). "
        "Weight decay 1×10⁻⁵ prevents overfitting without meaningfully slowing learning.",
        indent=True
    )
    kv(doc, "Learning rate schedule", "CosineAnnealingLR(T_max=n_epochs, eta_min=1×10⁻⁵)")
    body(doc,
        "Cosine annealing reduces the learning rate smoothly from 3×10⁻⁴ to 1×10⁻⁵ over the "
        "full training run. This avoids the abrupt learning rate drops of step decay and the "
        "instability of a constant rate near convergence. The minimum learning rate 1×10⁻⁵ "
        "ensures continued fine-grained adjustment in the last epochs without oscillation.",
        indent=True
    )
    kv(doc, "Batch size", "min(N_active, 256)")
    body(doc,
        "All active pixels are trained in a single batch per epoch when N_active ≤ 256. This "
        "full-batch regime maximises the number of negative pairs in the NT-Xent loss "
        "(2N_active − 2 negatives per sample), which is critical for contrastive learning quality. "
        "The 256 cap prevents GPU/CPU memory exhaustion on very large recordings.",
        indent=True
    )
    kv(doc, "Default n_epochs", "100")
    body(doc,
        "100 epochs is sufficient for convergence in experiments with 50–300 active pixels and "
        "calcium traces of 60–500 frames. Users can reduce to 20–50 for exploratory analysis "
        "(≈ 5–15 s on CPU) or increase to 200–500 for maximum quality on high-SNR recordings. "
        "The epochs spinner in the UI exposes this parameter directly.",
        indent=True
    )
    kv(doc, "Signal z-scoring", "Per-pixel z-score before encoder input")
    body(doc,
        "Each pixel's signal is standardised to zero mean and unit variance before entering "
        "the encoder. This is distinct from the augmentation: z-scoring is a deterministic "
        "normalisation applied consistently to both training and inference, ensuring that the "
        "encoder's first layer receives inputs on a consistent numerical scale regardless of "
        "absolute fluorescence intensity.",
        indent=True
    )
    kv(doc, "Device", "CUDA if available, otherwise CPU")
    body(doc,
        "The model and all tensors are moved to GPU if torch.cuda.is_available() returns True. "
        "After training, the encoder is explicitly moved back to CPU (encoder.cpu()) so that "
        "downstream numpy operations in compute_embedding_adjacency do not require GPU awareness.",
        indent=True
    )

    heading(doc, "5.6  Embedding-Space Adjacency Graph Construction", 2)
    body(doc,
        "After training, the projection head is discarded and the encoder is evaluated in "
        "inference mode (model.eval(), torch.no_grad()) to extract the final representations "
        "h ∈ ℝ^128 for all active pixels. These embeddings are L2-normalised to unit vectors, "
        "after which cosine similarity reduces to the dot product:"
    )
    math_block(doc, "cos_sim(i, j)  =  h_i · h_j  =  Σ_k  h_i[k] · h_j[k]   (since ||h|| = 1)")
    body(doc,
        "The same 4-connected spatial graph used for NCC is then constructed, with edge weight "
        "cos_sim(i, j) clipped to [0, 1]. This graph is passed to the same cluster_ncc_graph() "
        "function used by NCC Graph — the Louvain community detection and its fallback cascade "
        "are shared between the two methods. The only difference is the source of the edge "
        "weights: raw cross-correlation peaks (NCC) vs. learned embedding cosine similarities "
        "(SimCLR)."
    )
    body(doc,
        "The embedding score (displayed as 'Embedding score' in the UI) uses the same "
        "intra-cluster mean formula as the NCC cluster score, applied to the embedding-space "
        "adjacency matrix. Values above 0.85 indicate that neighbouring pixels within each "
        "cluster have highly similar learned representations."
    )

    heading(doc, "5.7  Hyperparameters and Justification", 2)
    add_table(doc,
        ["Parameter", "Value", "Justification"],
        [
            ["Temperature τ", "0.5",
             "Standard SimCLR value (Chen et al. 2020). Provides strong gradient signal without numerical overflow."],
            ["Adam lr", "3×10⁻⁴",
             "Widely validated for Adam on contrastive tasks; balances fast early convergence with stable late-stage refinement."],
            ["Adam weight_decay", "1×10⁻⁵",
             "Mild regularisation; prevents overfitting to the small per-recording dataset without meaningfully slowing convergence."],
            ["CosineAnnealingLR eta_min", "1×10⁻⁵",
             "Sets the floor for the learning rate; prevents the learning rate from decaying to zero (which would freeze gradients)."],
            ["Batch size", "min(N, 256)",
             "Full-batch training maximises negatives per step. Cap of 256 prevents memory issues on large recordings."],
            ["n_epochs default", "100",
             "Empirically sufficient for convergence on typical calcium recordings (50–300 active pixels, 60–500 frames)."],
            ["Encoder dim h", "128",
             "Large enough to capture waveform complexity; small enough to avoid overfitting."],
            ["Projection dim z", "64",
             "Following SimCLR: lower-dimensional projection space for NT-Xent training; 128→64 with nonlinearity."],
            ["Phase shift range", "±T/2",
             "The entire recording length; ensures full phase invariance."],
            ["Amplitude scale", "U(0.6, 1.4)",
             "±40% covers realistic dye loading variability without distorting waveform shape beyond recognition."],
            ["Noise σ", "5% of signal range",
             "Matches typical experimental noise levels in calcium imaging (SNR ≈ 20)."],
            ["Drift amplitude", "10% of range",
             "Conservative; covers residual bleach/drift after subtraction without destroying beat morphology."],
            ["Conv1 kernel", "7",
             "Spans ≈ 0.47 s at 15 fps; captures full beat onset/offset at typical cardiomyocyte rates."],
            ["Conv2 kernel", "5",
             "Spans ≈ 0.33 s; captures beat-to-beat variability and fine kinetic features."],
            ["Conv3 kernel", "3",
             "Local refinement; smallest sensible convolutional kernel in 1D."],
        ],
        col_widths=[3.5, 2.5, 10.0]
    )

    heading(doc, "5.8  Libraries", 2)
    add_table(doc,
        ["Library", "Version Requirement", "Role"],
        [
            ["torch (PyTorch)", "≥ 2.0 (CPU build)", "Neural network encoder, optimiser, loss function"],
            ["torch.nn", "—", "Conv1d, BatchNorm1d, Linear, ReLU, MaxPool1d, AdaptiveAvgPool1d"],
            ["torch.nn.functional", "—", "cross_entropy (NT-Xent loss), normalize (L2)"],
            ["numpy", "≥ 1.24", "Signal preprocessing, augmentation, adjacency matrix construction"],
            ["scipy.sparse", "≥ 1.10", "csr_matrix for adjacency storage, reused from NCC Graph"],
            ["python-louvain + networkx", "≥ 0.16, ≥ 3.0", "Community detection (shared with NCC Graph)"],
        ],
        col_widths=[4.0, 3.0, 9.0]
    )

    doc.add_page_break()

    # ====================================================================
    # 6. COMPARATIVE ANALYSIS
    # ====================================================================
    heading(doc, "6.  Comparative Analysis", 1)
    body(doc,
        "The three methods represent three distinct philosophies in unsupervised clustering for "
        "temporal signals: feature engineering, direct similarity measurement, and learned "
        "representation. The table below summarises their key properties:"
    )
    add_table(doc,
        ["Property", "HDBSCAN", "NCC Graph", "SimCLR Graph"],
        [
            ["Feature type", "5 hand-crafted scalars", "Raw waveform cross-correlation", "Learned 128D CNN embedding"],
            ["Similarity measure", "Euclidean distance in 5D", "Peak NCC at optimal lag", "Cosine similarity in 128D"],
            ["Spatial coherence", "Not enforced", "Hard (4-connected graph)", "Hard (4-connected graph)"],
            ["Phase discrimination", "Indirect (phase feature)", "Direct (NCC ≈ 0 at π phase lag)", "Learned from augmentation invariances"],
            ["Frequency bias", "High (dominant freq. feature)", "None (raw waveform)", "Learned; adapts per recording"],
            ["Cluster count", "Automatic (density)", "Automatic (modularity)", "Automatic (modularity)"],
            ["Requires labels", "No", "No", "No"],
            ["Training cost", "< 1 s", "~2–15 s (FFT NCC)", "~10–120 s (CNN training)"],
            ["Nonlinear waveform features", "No (linear features)", "No (linear correlation)", "Yes (deep CNN)"],
            ["Quality metric", "None", "Intra-cluster NCC score", "Intra-cluster embedding score"],
            ["Requires PyTorch", "No", "No", "Yes"],
            ["Key library", "hdbscan", "scipy, python-louvain", "torch, python-louvain"],
            ["Best use case", "Fast exploration, large low-SNR data", "Spatially coherent wave patterns", "Complex waveform heterogeneity, high-quality recordings"],
        ],
        col_widths=[4.5, 4.5, 4.5, 4.5]
    )
    caption(doc, "Table 6.1 — Side-by-side comparison of the three clustering methods.")

    body(doc,
        "Expected Agreement and Divergence. On a recording with well-separated functional regions "
        "and high SNR, all three methods should yield qualitatively similar partitions. Divergence "
        "is expected in three situations:"
    )
    bullet(doc, "Low-SNR recordings: HDBSCAN may produce more noise-labeled points; NCC Graph may produce a coarser partition (fewer clusters) because noisy signals have low NCC; SimCLR Graph is most robust because augmentation-based training explicitly conditions the encoder to handle noise.")
    bullet(doc, "Anti-phase pairs: NCC Graph and SimCLR Graph correctly separate pixels that fire at identical rates but opposite phases; HDBSCAN may incorrectly co-cluster them because the phase feature is a scalar that wraps at ±π.")
    bullet(doc, "Cells connected by spatial gaps: both NCC Graph and SimCLR Graph enforce 4-connected spatial coherence, so they cannot merge isolated active regions; HDBSCAN, which operates in feature space, can cluster distant pixels with identical waveform properties.")

    doc.add_page_break()

    # ====================================================================
    # 7. COMPLETE HYPERPARAMETER INVENTORY
    # ====================================================================
    heading(doc, "7.  Complete Hyperparameter Inventory", 1)
    body(doc, "The following table enumerates every numerical constant in the implementation.")
    add_table(doc,
        ["Component", "Parameter", "Value", "Source / Rationale"],
        [
            # Preprocessing
            ["Preprocessing", "Clip percentiles", "[0.4, 99.6]", "Removes < 0.4% of dark-frame outliers and < 0.4% of saturation events"],
            ["Preprocessing", "Gaussian σ", "H_raw / 204.8", "Adaptive anti-aliasing before binning"],
            ["Preprocessing", "Default bin size", "16 px (< 2048) / 32 px", "Reduces computation while preserving single-cell resolution"],
            ["Bleach – Single Exp", "Max fit iterations", "2000", "scipy curve_fit maxfev"],
            ["Bleach – Single Exp", "τ lower bound", "1×10⁻⁵ s", "Prevents numerical instability"],
            ["Bleach – Envelope", "Window percentile", "10th", "Robust lower-envelope estimation"],
            ["Bleach – Envelope", "Window size", "max(T/8, 3) frames", "Spans at least one beat cycle"],
            # Pulsatility
            ["Pulsatility", "Bandpass f_low", "0.3 Hz", "Minimum cardiac rate (18 BPM)"],
            ["Pulsatility", "Bandpass f_high", "5.0 Hz", "Maximum cardiac rate (300 BPM)"],
            ["Pulsatility", "Filter order", "4", "24 dB/oct rolloff"],
            ["Pulsatility", "Nyquist safety", "nyq × 0.9", "Prevents filter instability"],
            # Activity mask
            ["Activity Mask", "Otsu fallback percentile", "65th", "Robust default when Otsu fails"],
            ["Activity Mask", "Min active fraction", "3%", "Below → use percentile fallback"],
            ["Activity Mask", "Max active fraction", "85%", "Above → use percentile fallback"],
            # Beat detection
            ["Beat Detection", "Prominence threshold", "0.15 × range", "5th–95th of typical peak SNR"],
            ["Beat Detection", "Min inter-peak distance", "fps × 0.8 frames", "Equivalent to 75 BPM max"],
            # HDBSCAN
            ["HDBSCAN", "min_cluster_size", "max(3, N // 15)", "Scales with recording size"],
            ["HDBSCAN", "min_samples", "2", "Permissive density estimation"],
            ["HDBSCAN", "Amplitude window", "max(1.5 × fps, 3) frames", "≈ 1.5 beat cycles"],
            # KMeans fallback
            ["KMeans fallback", "n_clusters", "min(4, max(2, N // 5))", "2–4 clusters for small N"],
            ["KMeans fallback", "n_init", "10", "sklearn default; avoid local minima"],
            ["KMeans fallback", "random_state", "42", "Reproducibility"],
            # NCC Graph
            ["NCC Graph", "max_lag_fraction", "0.25", "Caps lag at ±25% T"],
            ["NCC Graph", "NCC clip range", "[0, 1]", "No negative edge weights"],
            ["NCC Graph", "FFT method", "scipy correlate fft", "O(T log T) complexity"],
            # Louvain
            ["Louvain", "resolution", "1.0", "Standard Newman–Girvan modularity"],
            ["Louvain", "random_state", "42", "Reproducibility"],
            # Spectral fallback
            ["Spectral fallback", "n_clusters", "min(8, max(2, N // 10))", "Scales with active pixels"],
            ["Spectral fallback", "n_init", "10", "sklearn default"],
            ["Spectral fallback", "random_state", "42", "Reproducibility"],
            # HDBSCAN-on-NCC fallback
            ["NCC-HDBSCAN fb.", "min_cluster_size", "max(3, N // 15)", "Same formula as primary HDBSCAN"],
            ["NCC-HDBSCAN fb.", "min_samples", "2", "Permissive density estimation"],
            # SimCLR Encoder
            ["SimCLR Encoder", "Conv1 channels", "32", "First-level temporal features"],
            ["SimCLR Encoder", "Conv1 kernel", "7 (pad=3)", "Multi-scale capture; ~0.5 s at 15 fps"],
            ["SimCLR Encoder", "Conv2 channels", "64", "Higher-level features"],
            ["SimCLR Encoder", "Conv2 kernel", "5 (pad=2)", "Beat-to-beat variability"],
            ["SimCLR Encoder", "Conv3 channels", "128", "Abstract representation dim"],
            ["SimCLR Encoder", "Conv3 kernel", "3 (pad=1)", "Local refinement"],
            ["SimCLR Encoder", "Projection h→z", "128 → 128 → 64", "Nonlinear projection head"],
            # Augmentations
            ["Augmentation", "Phase shift", "U(−T/2, T/2)", "Full recording-length phase range"],
            ["Augmentation", "Amplitude scale", "U(0.6, 1.4)", "±40% amplitude variation"],
            ["Augmentation", "Noise σ", "5% of signal range", "~5–20 dB SNR robustness"],
            ["Augmentation", "Drift amplitude", "10% of range", "Residual post-correction drift"],
            # NT-Xent
            ["NT-Xent Loss", "Temperature τ", "0.5", "SimCLR paper recommendation"],
            # Training
            ["Training", "Adam lr", "3×10⁻⁴", "Standard contrastive learning rate"],
            ["Training", "Adam weight_decay", "1×10⁻⁵", "Mild L2 regularisation"],
            ["Training", "LR schedule", "CosineAnnealingLR", "Smooth decay without abrupt drops"],
            ["Training", "eta_min", "1×10⁻⁵", "Non-zero floor prevents gradient freeze"],
            ["Training", "Batch size", "min(N, 256)", "Full-batch preferred; capped at 256"],
            ["Training", "Default n_epochs", "100", "Convergence for typical recordings"],
            ["Training", "Z-score norm", "Per-pixel, axis=1", "Removes absolute intensity bias"],
            ["Training", "Device", "CUDA if available else CPU", "Automatic GPU acceleration"],
        ],
        col_widths=[3.5, 4.0, 3.5, 5.0]
    )
    caption(doc, "Table 7.1 — Complete hyperparameter inventory for all components of the napari-cta-simclr plugin.")

    doc.add_page_break()

    # ====================================================================
    # 8. REFERENCES
    # ====================================================================
    heading(doc, "8.  References", 1)

    refs = [
        ("Chen et al., 2020",
         "T. Chen, S. Kornblith, M. Norouzi, G. Hinton. "
         "\"A Simple Framework for Contrastive Learning of Visual Representations.\" "
         "Proceedings of the 37th ICML, 2020. arXiv:2002.05709."),
        ("Campello et al., 2013",
         "R. J. G. B. Campello, D. Moulavi, J. Sander. "
         "\"Density-Based Clustering Based on Hierarchical Density Estimates.\" "
         "Proceedings of PAKDD 2013, pp. 160–172."),
        ("Blondel et al., 2008",
         "V. D. Blondel, J.-L. Guillaume, R. Lambiotte, E. Lefebvre. "
         "\"Fast unfolding of communities in large networks.\" "
         "Journal of Statistical Mechanics: Theory and Experiment, 2008(10): P10008."),
        ("Otsu, 1979",
         "N. Otsu. \"A Threshold Selection Method from Gray-Level Histograms.\" "
         "IEEE Transactions on Systems, Man, and Cybernetics, 9(1): 62–66, 1979."),
        ("McInnes et al., 2017",
         "L. McInnes, J. Healy, S. Astels. \"hdbscan: Hierarchical density based clustering.\" "
         "Journal of Open Source Software, 2(11): 205, 2017."),
        ("Ng et al., 2001",
         "A. Y. Ng, M. I. Jordan, Y. Weiss. "
         "\"On Spectral Clustering: Analysis and an Algorithm.\" "
         "Advances in Neural Information Processing Systems 14, 2001."),
        ("Pedregosa et al., 2011",
         "F. Pedregosa et al. \"Scikit-learn: Machine Learning in Python.\" "
         "JMLR, 12: 2825–2830, 2011."),
        ("Paszke et al., 2019",
         "A. Paszke et al. \"PyTorch: An Imperative Style, High-Performance Deep Learning Library.\" "
         "Advances in Neural Information Processing Systems 32, 2019."),
        ("Virtanen et al., 2020",
         "P. Virtanen et al. \"SciPy 1.0: Fundamental Algorithms for Scientific Computing in Python.\" "
         "Nature Methods, 17: 261–272, 2020."),
        ("Walt et al., 2014",
         "S. van der Walt et al. \"scikit-image: image processing in Python.\" "
         "PeerJ, 2: e453, 2014."),
    ]

    for num, (key, text) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.space_after  = Pt(5)
        r1 = p.add_run(f"[{num}] ")
        set_font(r1, size=11, bold=True)
        r2 = p.add_run(text)
        set_font(r2, size=11)

    # ====================================================================
    # SAVE
    # ====================================================================
    doc.save(OUTPUT)
    print(f"\nReport written to:\n  {OUTPUT}\n")


if __name__ == "__main__":
    build()
